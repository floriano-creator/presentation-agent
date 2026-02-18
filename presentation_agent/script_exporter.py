"""Export speech manuscript to Word (.docx)."""

from __future__ import annotations

from pathlib import Path
from typing import Optional

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from presentation_agent.models import PresentationManuscript


class ScriptExporter:
    """Exports the full speech manuscript to a Word document."""

    def export(
        self,
        manuscript: PresentationManuscript,
        output_path: str | Path,
        topic: Optional[str] = None,
        duration_minutes: Optional[int] = None,
        audience: Optional[str] = None,
    ) -> str:
        """
        Export manuscript to a .docx file.

        The document contains the full spoken script: title, metadata,
        and continuous prose—suitable for reading aloud.

        Args:
            manuscript: The presentation manuscript.
            output_path: Path for the .docx file.
            topic: Optional topic for metadata.
            duration_minutes: Optional duration for metadata.
            audience: Optional audience for metadata.

        Returns:
            Resolved path to the exported file.
        """
        path = Path(output_path)
        path.parent.mkdir(parents=True, exist_ok=True)

        doc = Document()

        # Title (centered)
        title = doc.add_heading(manuscript.title, level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()

        # Optional metadata
        meta_parts = []
        if topic:
            meta_parts.append(f"Topic: {topic}")
        if duration_minutes is not None:
            meta_parts.append(f"Duration: {duration_minutes} min")
        if audience:
            meta_parts.append(f"Audience: {audience}")

        if meta_parts:
            meta = doc.add_paragraph(" | ".join(meta_parts))
            meta.paragraph_format.space_after = Pt(12)

        doc.add_paragraph()

        # Full manuscript content: continuous prose
        for section in manuscript.sections:
            # Section heading (optional, for structure)
            if section.name:
                doc.add_heading(section.name, level=1)

            # Section content as clean paragraphs
            # Split by double newlines for paragraph breaks
            paragraphs = [
                p.strip() for p in section.content.split("\n\n") if p.strip()
            ]
            for para_text in paragraphs:
                p = doc.add_paragraph(para_text)
                p.paragraph_format.space_after = Pt(6)

            doc.add_paragraph()

        doc.save(str(path))
        return str(path.resolve())
